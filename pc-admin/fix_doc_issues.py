#!/usr/bin/env python3
"""
修复响应文件的三个问题：
1. 表格中的<br>换行
2. 补充售后服务方案大纲
3. 补充质量保障方案大纲
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

TEMPLATE_PATH = '/workspace/.uploads/8de091e5-b565-412c-8aa2-e4702d5e80f5_00f9f188-b18e-476a-b731-854c69494b4c_响应文件 - V0.1.docx'

def set_font(run, size=12, name='宋体', color=None):
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_soft_break(run):
    """在run中添加软回车（Shift+Enter）"""
    br = etree.SubElement(run._element, qn('w:br'))

def fix_br_in_cell(cell):
    """修复单元格中的<br>文本为软回车"""
    for para in cell.paragraphs:
        text = para.text
        if text and '<br>' in text:
            parts = text.split('<br>')
            # 清除原段落内容
            para.clear()
            for i, part in enumerate(parts):
                part = part.strip()
                if part:
                    run = para.add_run(part)
                    set_font(run)
                if i < len(parts) - 1:
                    # 添加软回车
                    run = para.add_run('')
                    set_font(run)
                    add_soft_break(run)

def fix_br_in_tables(doc):
    """修复所有表格中的<br>"""
    fixed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '<br>' in para.text:
                        fix_br_in_cell(cell)
                        fixed += 1
                        break  # 一个单元格只修复一次
    print(f"Fixed <br> in {fixed} table cells")
    return fixed

def add_heading(doc, text, level):
    """添加标题并设置字体"""
    p = doc.add_paragraph(text, style=f'Heading {level}')
    for run in p.runs:
        set_font(run)
    return p

def add_paragraph(doc, text, bold=False):
    """添加段落并设置字体"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run)
    if bold:
        run.bold = True
    return p

def add_after_heading(doc, target_text, new_elements):
    """在指定标题后插入新元素"""
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if target_text in p.text and p.style.name.startswith('Heading'):
            target_idx = i
            break
    
    if target_idx is None:
        print(f"WARNING: Could not find heading '{target_text}'")
        return False
    
    # 找到该标题下的最后一个段落（直到下一个同级或更高级标题）
    insert_idx = target_idx + 1
    for i in range(target_idx + 1, len(doc.paragraphs)):
        style = doc.paragraphs[i].style.name
        if style.startswith('Heading'):
            level = int(style.replace('Heading ', '')) if style.replace('Heading ', '').isdigit() else 99
            target_level = int(doc.paragraphs[target_idx].style.name.replace('Heading ', '')) if doc.paragraphs[target_idx].style.name.replace('Heading ', '').isdigit() else 99
            if level <= target_level:
                break
        insert_idx = i + 1
    
    # 在insert_idx位置插入新元素
    body = doc.element.body
    for elem in new_elements:
        body.insert(insert_idx, elem)
        insert_idx += 1
    
    return True

def main():
    print("Loading document...")
    doc = Document(TEMPLATE_PATH)
    
    # ===== 1. 修复表格中的<br> =====
    print("\n=== Fixing <br> in tables ===")
    fix_br_in_tables(doc)
    
    # 保存并重新加载
    doc.save(TEMPLATE_PATH)
    doc = Document(TEMPLATE_PATH)
    
    # ===== 2 & 3. 补充售后服务方案和质量保障方案 =====
    print("\n=== Adding 售后服务方案 and 质量保障方案 ===")
    
    # 在"第四部分 项目组织与保障"后面插入
    # 先找到这个位置
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if '第四部分 项目组织与保障' in p.text and p.style.name == 'Heading 2':
            target_idx = i
            print(f"Found target at [{i}]: {p.text.strip()[:40]}")
            break
    
    if target_idx is None:
        print("WARNING: Could not find '第四部分 项目组织与保障', appending to end")
        # 找到文档末尾
        target_idx = len(doc.paragraphs) - 1
    
    # 找到插入位置（下一个Heading 2之前或文档末尾）
    insert_pos = target_idx + 1
    for i in range(target_idx + 1, len(doc.paragraphs)):
        style = doc.paragraphs[i].style.name
        if style == 'Heading 2':
            insert_pos = i
            break
        insert_pos = i + 1
    
    body = doc.element.body
    
    # 创建售后服务方案内容
    aftersales_elements = []
    
    # Heading 2: 售后服务方案
    p = doc.add_paragraph('售后服务方案', style='Heading 2')
    for run in p.runs:
        set_font(run)
    aftersales_elements.append(p._element)
    
    # Heading 3: 售后服务体系
    p = doc.add_paragraph('售后服务体系', style='Heading 3')
    for run in p.runs:
        set_font(run)
    aftersales_elements.append(p._element)
    
    p = doc.add_paragraph('建立完善的售后服务体系，包括技术支持团队、运维服务团队和客户成功团队。设立7×24小时服务热线，提供电话、邮件、微信等多种服务渠道，确保用户在遇到问题时能够及时获得帮助。')
    for run in p.runs:
        set_font(run)
    aftersales_elements.append(p._element)
    
    # Heading 3: 售后服务内容
    p = doc.add_paragraph('售后服务内容', style='Heading 3')
    for run in p.runs:
        set_font(run)
    aftersales_elements.append(p._element)
    
    p = doc.add_paragraph('售后服务内容包括：系统故障排查与修复、性能优化建议、功能使用指导、数据备份与恢复支持、安全漏洞修复、系统升级服务等。对于紧急故障，承诺2小时内响应，24小时内解决；一般问题承诺4小时内响应，48小时内解决。')
    for run in p.runs:
        set_font(run)
    aftersales_elements.append(p._element)
    
    # Heading 3: 培训计划
    p = doc.add_paragraph('培训计划', style='Heading 3')
    for run in p.runs:
        set_font(run)
    aftersales_elements.append(p._element)
    
    p = doc.add_paragraph('提供系统化的培训计划，包括管理员培训、运营人员培训和终端用户培训。培训方式包括现场培训、线上视频培训和操作手册自学。培训内容涵盖系统功能操作、日常运维管理、常见问题处理等，确保用户能够熟练使用系统。')
    for run in p.runs:
        set_font(run)
    aftersales_elements.append(p._element)
    
    # Heading 3: 维保期后的运维服务内容
    p = doc.add_paragraph('维保期后的运维服务内容', style='Heading 3')
    for run in p.runs:
        set_font(run)
    aftersales_elements.append(p._element)
    
    p = doc.add_paragraph('维保期结束后，提供有偿运维服务延续方案。服务内容包括：系统日常巡检、版本更新升级、数据迁移支持、功能迭代开发、技术咨询服务等。根据用户需求提供年度运维服务合同，保障系统长期稳定运行。')
    for run in p.runs:
        set_font(run)
    aftersales_elements.append(p._element)
    
    # 创建质量保障方案内容
    quality_elements = []
    
    # Heading 2: 质量保障方案
    p = doc.add_paragraph('质量保障方案', style='Heading 2')
    for run in p.runs:
        set_font(run)
    quality_elements.append(p._element)
    
    # Heading 3: 质量保障体系
    p = doc.add_paragraph('质量保障体系', style='Heading 3')
    for run in p.runs:
        set_font(run)
    quality_elements.append(p._element)
    
    p = doc.add_paragraph('建立完善的质量保障体系，覆盖需求分析、设计、开发、测试、交付全流程。设立质量管理小组，制定质量标准和检查清单，实施代码审查、单元测试、集成测试、系统测试和用户验收测试等多层质量关卡。')
    for run in p.runs:
        set_font(run)
    quality_elements.append(p._element)
    
    # Heading 3: 质量保障措施
    p = doc.add_paragraph('质量保障措施', style='Heading 3')
    for run in p.runs:
        set_font(run)
    quality_elements.append(p._element)
    
    p = doc.add_paragraph('质量保障措施包括：严格执行代码规范和开发标准；实施持续集成和自动化测试；定期进行代码审查和技术评审；建立缺陷跟踪和管理机制；开展用户验收测试和试运行验证；提供完善的系统文档和操作手册。')
    for run in p.runs:
        set_font(run)
    quality_elements.append(p._element)
    
    # Heading 3: 质量风险控制
    p = doc.add_paragraph('质量风险控制', style='Heading 3')
    for run in p.runs:
        set_font(run)
    quality_elements.append(p._element)
    
    p = doc.add_paragraph('建立质量风险识别和防控机制，对项目全生命周期中的质量风险进行识别、评估和管控。主要风险包括：需求变更风险、技术实现风险、进度延期风险、人员变动风险等。制定风险应对预案，建立风险监控和预警机制，确保项目质量目标达成。')
    for run in p.runs:
        set_font(run)
    quality_elements.append(p._element)
    
    # Heading 3: 应急服务机制
    p = doc.add_paragraph('应急服务机制', style='Heading 3')
    for run in p.runs:
        set_font(run)
    quality_elements.append(p._element)
    
    p = doc.add_paragraph('建立完善的应急服务机制，包括应急响应流程、应急预案和应急资源储备。设立应急响应小组，明确应急响应等级和处理流程。对于系统重大故障，启动应急预案，确保在最短时间内恢复系统正常运行。定期进行应急演练，提升团队应急处理能力。')
    for run in p.runs:
        set_font(run)
    quality_elements.append(p._element)
    
    # 插入到文档中
    # 先插入售后服务方案，再插入质量保障方案
    all_elements = aftersales_elements + quality_elements
    
    for elem in reversed(all_elements):
        body.insert(insert_pos, elem)
    
    print(f"Inserted {len(all_elements)} elements at position {insert_pos}")
    
    # 保存
    doc.save(TEMPLATE_PATH)
    print("\nSaved!")
    
    # 验证
    doc2 = Document(TEMPLATE_PATH)
    print(f"\nFinal: {len(doc2.paragraphs)} paragraphs, {len(doc2.tables)} tables")
    
    # 检查<br>残留
    br_remaining = 0
    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '<br>' in para.text:
                        br_remaining += 1
    print(f"Remaining <br> in tables: {br_remaining}")
    
    # 检查新章节
    print("\n=== New chapter structure ===")
    for i, p in enumerate(doc2.paragraphs):
        style = p.style.name if p.style else ''
        if style.startswith('Heading') and any(k in p.text for k in ['售后','质量保障']):
            print(f'[{i}] {style:<12} {p.text.strip()[:50]}')

if __name__ == '__main__':
    main()
