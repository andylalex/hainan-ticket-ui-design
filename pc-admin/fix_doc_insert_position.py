#!/usr/bin/env python3
"""
修复章节插入位置错误。
删除错误插入的售后服务方案和质量保障方案，然后在正确位置重新插入。
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE_PATH = '/workspace/.uploads/8de091e5-b565-412c-8aa2-e4702d5e80f5_00f9f188-b18e-476a-b731-854c69494b4c_响应文件 - V0.1.docx'

def set_font(run, size=12, name='宋体'):
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def main():
    print("Loading document...")
    doc = Document(TEMPLATE_PATH)
    body = doc.element.body

    # ===== 1. 删除错误插入的段落 [4911] 到 [4928] =====
    # 先找到所有"售后服务方案"和"质量保障方案"相关的段落
    remove_indices = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ''
        if style.startswith('Heading') and ('售后服务方案' in text or '质量保障方案' in text):
            remove_indices.append(i)
            print(f"Found heading to remove/reposition: [{i}] {style} {text[:40]}")
        elif style == 'Heading 3' and text in ('售后服务体系', '售后服务内容', '培训计划', '维保期后的运维服务内容',
                                               '质量保障体系', '质量保障措施', '质量风险控制', '应急服务机制'):
            remove_indices.append(i)
            print(f"Found subheading: [{i}] {style} {text[:40]}")

    # 同时找到这些标题下面的正文段落
    all_remove_indices = set(remove_indices)
    for idx in remove_indices:
        # 检查下一个段落是否是Normal（正文）
        if idx + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[idx + 1]
            if next_p.style.name == 'Normal' and next_p.text.strip():
                all_remove_indices.add(idx + 1)
                print(f"Found body text: [{idx+1}] {next_p.text[:40]}")

    print(f"\nRemoving {len(all_remove_indices)} paragraphs...")
    for idx in sorted(all_remove_indices, reverse=True):
        try:
            body.remove(doc.paragraphs[idx]._element)
        except (ValueError, IndexError) as e:
            print(f"  Warning: could not remove [{idx}]: {e}")

    # 保存并重新加载
    doc.save(TEMPLATE_PATH)
    doc = Document(TEMPLATE_PATH)
    body = doc.element.body
    print(f"After removal: {len(doc.paragraphs)} paragraphs")

    # ===== 2. 在正确位置重新插入 =====
    # 找到"第四部分 项目组织与保障"的位置
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if '第四部分 项目组织与保障' in p.text and p.style.name == 'Heading 2':
            target_idx = i
            print(f"\nFound target at [{i}]: {p.text.strip()[:40]}")
            break

    if target_idx is None:
        print("ERROR: Could not find target heading")
        return

    # 找到插入位置：这个Heading 2后面所有的内容，直到文档末尾或下一个Heading 2
    # 由于"第四部分"是文档末尾的章节，我们在它后面所有内容的末尾插入
    # 找到"第四部分"章节内容的末尾
    insert_after = target_idx
    for i in range(target_idx + 1, len(doc.paragraphs)):
        style = doc.paragraphs[i].style.name
        if style == 'Heading 2':
            break
        insert_after = i

    print(f"Insert after paragraph [{insert_after}]")
    target_elem = doc.paragraphs[insert_after]._element

    # 创建新内容
    new_elements = []

    def create_heading(text, level):
        p = doc.add_paragraph(text, style=f'Heading {level}')
        for run in p.runs:
            set_font(run)
        return p._element

    def create_body(text):
        p = doc.add_paragraph(text)
        for run in p.runs:
            set_font(run)
        return p._element

    # 售后服务方案
    new_elements.append(create_heading('售后服务方案', 2))
    new_elements.append(create_heading('售后服务体系', 3))
    new_elements.append(create_body('建立完善的售后服务体系，包括技术支持团队、运维服务团队和客户成功团队。设立7×24小时服务热线，提供电话、邮件、微信等多种服务渠道，确保用户在遇到问题时能够及时获得帮助。'))
    new_elements.append(create_heading('售后服务内容', 3))
    new_elements.append(create_body('售后服务内容包括：系统故障排查与修复、性能优化建议、功能使用指导、数据备份与恢复支持、安全漏洞修复、系统升级服务等。对于紧急故障，承诺2小时内响应，24小时内解决；一般问题承诺4小时内响应，48小时内解决。'))
    new_elements.append(create_heading('培训计划', 3))
    new_elements.append(create_body('提供系统化的培训计划，包括管理员培训、运营人员培训和终端用户培训。培训方式包括现场培训、线上视频培训和操作手册自学。培训内容涵盖系统功能操作、日常运维管理、常见问题处理等，确保用户能够熟练使用系统。'))
    new_elements.append(create_heading('维保期后的运维服务内容', 3))
    new_elements.append(create_body('维保期结束后，提供有偿运维服务延续方案。服务内容包括：系统日常巡检、版本更新升级、数据迁移支持、功能迭代开发、技术咨询服务等。根据用户需求提供年度运维服务合同，保障系统长期稳定运行。'))

    # 质量保障方案
    new_elements.append(create_heading('质量保障方案', 2))
    new_elements.append(create_heading('质量保障体系', 3))
    new_elements.append(create_body('建立完善的质量保障体系，覆盖需求分析、设计、开发、测试、交付全流程。设立质量管理小组，制定质量标准和检查清单，实施代码审查、单元测试、集成测试、系统测试和用户验收测试等多层质量关卡。'))
    new_elements.append(create_heading('质量保障措施', 3))
    new_elements.append(create_body('质量保障措施包括：严格执行代码规范和开发标准；实施持续集成和自动化测试；定期进行代码审查和技术评审；建立缺陷跟踪和管理机制；开展用户验收测试和试运行验证；提供完善的系统文档和操作手册。'))
    new_elements.append(create_heading('质量风险控制', 3))
    new_elements.append(create_body('建立质量风险识别和防控机制，对项目全生命周期中的质量风险进行识别、评估和管控。主要风险包括：需求变更风险、技术实现风险、进度延期风险、人员变动风险等。制定风险应对预案，建立风险监控和预警机制，确保项目质量目标达成。'))
    new_elements.append(create_heading('应急服务机制', 3))
    new_elements.append(create_body('建立完善的应急服务机制，包括应急响应流程、应急预案和应急资源储备。设立应急响应小组，明确应急响应等级和处理流程。对于系统重大故障，启动应急预案，确保在最短时间内恢复系统正常运行。定期进行应急演练，提升团队应急处理能力。'))

    # 在target_elem后插入所有新元素
    current = target_elem
    for elem in new_elements:
        current.addnext(elem)
        current = elem

    print(f"Inserted {len(new_elements)} new elements")

    # 保存
    doc.save(TEMPLATE_PATH)
    print("Saved!")

    # 验证
    doc2 = Document(TEMPLATE_PATH)
    print(f"\nFinal: {len(doc2.paragraphs)} paragraphs")

    # 检查新章节位置
    print("\n=== New chapter positions ===")
    for i, p in enumerate(doc2.paragraphs):
        style = p.style.name if p.style else ''
        if style.startswith('Heading') and any(k in p.text for k in ['售后','质量保障','第四部分']):
            print(f'[{i}] {style:<12} {p.text.strip()[:50]}')

if __name__ == '__main__':
    main()
